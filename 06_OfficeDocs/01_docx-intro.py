import docx

d = docx.Document()
p = d.add_paragraph()
#p.add_run("Hello, this is a text, укр")
p.add_run("HELLO, this is a text, укр TEXT 123 hello")

d.save("file01.docx")
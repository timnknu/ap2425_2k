import docx

d = docx.Document()
p = d.add_paragraph()
r = p.add_run("HELLO, this is a text, укр TEXT 123 hello")
r.font.color.rgb = docx.shared.RGBColor(250, 0, 250)

d.save("file01.docx")